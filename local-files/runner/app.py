# ============================================================================
# RUNNER ALURA - FastAPI + Playwright + LibreOffice
# ============================================================================

import os
import json
import time
import re
import subprocess
import shutil
import uuid
import tempfile
from pathlib import Path
from typing import List, Optional
from io import BytesIO

# Third-party
from unidecode import unidecode
import unicodedata
from bs4 import BeautifulSoup, NavigableString
from urllib.parse import urlencode, urljoin
import requests
import httpx
from openai import OpenAI

# FastAPI
from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks, Form
from fastapi.responses import JSONResponse, Response, FileResponse
from pydantic import BaseModel, field_validator

# Playwright
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeout

# DOCX
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# PIL - DEVE ser importado ANTES do UNO para evitar conflito de imports
from PIL import Image as PILImage

# LibreOffice UNO (opcional)
try:
    import uno
    from com.sun.star.beans import PropertyValue
    LIBREOFFICE_DISPONIVEL = True
except ImportError:
    LIBREOFFICE_DISPONIVEL = False
    print("python3-uno nao disponivel - endpoints LibreOffice desabilitados")

# Track Changes OOXML
from track_changes import aplicar_revisoes_docx, aplicar_comentarios_docx, TrackChangesApplicator

# LLM e Prompts para agentes de revisao
from llm_client import criar_cliente_llm
from prompts_revisao import (
    formatar_prompt_seo,
    formatar_prompt_tecnico,
    formatar_prompt_texto,
    formatar_prompt_imagem
)


# ============================================================================
# CONFIGURAÇÃO
# ============================================================================

client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

TEMP_DIR = Path("/tmp/video_processing")
TEMP_DIR.mkdir(exist_ok=True)


# ============================================================================
# MODELOS PYDANTIC
# ============================================================================

# --- Vídeo ---
class LegendaConfig(BaseModel):
    font_size: int = 24
    font_name: str = "Arial"
    bold: bool = False
    primary_colour: str = "&HFFFFFF"
    outline_colour: str = "&H000000"
    back_colour: str = "&H80000000"
    outline: int = 2
    shadow: int = 1
    margin_v: int = 30


class VideoURLProcessingPayload(BaseModel):
    video_urls: List[str]
    audio_url: str
    transicao_duracao: float = 0.5
    transicao_tipo: str = "fade"
    output_filename: str = "video_final.mp4"
    adicionar_legendas: bool = False
    estilo_legenda: str = "youtube"
    legenda_config: Optional[LegendaConfig] = None


# --- LinkedIn/Alura ---
class PesquisaPayload(BaseModel):
    query: str
    n_vagas: int


class Payload(BaseModel):
    nome_curso: str
    nome_instrutor: str
    tempo_curso: int


class IDPayload(BaseModel):
    id: str


# --- DOCX ---
class TextSegment(BaseModel):
    text: str
    link: Optional[str] = None
    bold: Optional[bool] = False
    italic: Optional[bool] = False


class ContentItem(BaseModel):
    type: str
    level: Optional[int] = None
    text: Optional[str] = None
    segments: Optional[List[TextSegment]] = None
    ordered: Optional[bool] = False
    items: Optional[List] = None
    language: Optional[str] = None
    content: Optional[str] = None
    url: Optional[str] = None
    alt: Optional[str] = None
    width: Optional[int] = None
    height: Optional[int] = None
    headers: Optional[List[str]] = None
    rows: Optional[List[List[str]]] = None
    cite: Optional[str] = None

    @field_validator('segments', mode='before')
    @classmethod
    def filter_none_segments(cls, v):
        if v is None:
            return None
        return [seg for seg in v if seg is not None]

    @field_validator('items', mode='before')
    @classmethod
    def filter_none_items(cls, v):
        if v is None:
            return None
        return [item for item in v if item is not None]


class ArticleMetadata(BaseModel):
    title: Optional[str] = None
    author: Optional[str] = None
    publishDate: Optional[str] = None


class GenerateDocxPayload(BaseModel):
    metadata: ArticleMetadata
    content: List[ContentItem]
    filename: Optional[str] = "documento.docx"
    base_url: Optional[str] = None

    @field_validator('content', mode='before')
    @classmethod
    def filter_none_content(cls, v):
        if v is None:
            return []
        return [item for item in v if item is not None]


class ExtractArticlePayload(BaseModel):
    url: str


# --- LibreOffice ---
class RevisaoLibreOffice(BaseModel):
    """Uma revisão a ser aplicada ao documento."""
    tipo: str  # "SEO", "TECNICO", "TEXTO"
    acao: str  # "substituir", "deletar", "inserir", "comentario"
    paragrafo: int  # Índice do parágrafo (0-indexed)
    inicio: int  # Posição inicial no parágrafo
    fim: int  # Posição final no parágrafo
    texto_esperado: str = ""  # Texto que DEVERIA estar nesta posição (validação)
    texto_novo: Optional[str] = ""
    justificativa: str = ""


class ExtrairTextoResponse(BaseModel):
    """Resposta da extração de texto."""
    paragrafos: list
    texto_completo: str
    total_paragrafos: int


# --- Revisão com Track Changes (OOXML) ---
class RevisaoItem(BaseModel):
    """Uma revisão individual para aplicar ao documento."""
    tipo: str  # "SEO", "TECNICO", "TEXTO"
    acao: str  # "substituir", "deletar", "inserir", "comentario"
    texto_original: str  # Texto exato a ser encontrado no documento
    texto_novo: Optional[str] = ""  # Texto substituto (para substituir/inserir)
    justificativa: str = ""  # Explicação da mudança


class AplicarRevisoesPayload(BaseModel):
    """Payload para aplicar revisões via JSON."""
    docx_url: Optional[str] = None  # URL do documento DOCX
    docx_base64: Optional[str] = None  # Ou base64 do documento
    revisoes: List[RevisaoItem]  # Lista de revisões
    autor: str = "Agente IA Revisor"


class ExtrairTextoDocxPayload(BaseModel):
    """Payload para extrair texto de DOCX."""
    docx_url: Optional[str] = None
    docx_base64: Optional[str] = None


class RevisaoAgentPayload(BaseModel):
    """Payload para executar um agente de revisão."""
    docx_url: Optional[str] = None  # URL do documento DOCX
    docx_base64: Optional[str] = None  # Ou base64 do documento
    provider: str = "anthropic"  # "anthropic" ou "openai"
    guia_seo_url: Optional[str] = None  # URL do guia de SEO (apenas para agente SEO)
    url_artigo: Optional[str] = ""  # URL original do artigo (contexto)
    titulo: Optional[str] = ""  # Título do artigo (contexto)
    data_publicacao: Optional[str] = ""  # Data de publicação (contexto)


class RevisaoImagemPayload(BaseModel):
    """Payload para o agente de revisão de imagens."""
    docx_url: Optional[str] = None  # URL do documento DOCX
    docx_base64: Optional[str] = None  # Ou base64 do documento
    provider: str = "anthropic"  # "anthropic" ou "openai"
    url_artigo: str  # URL original do artigo (OBRIGATÓRIO para extrair imagens)
    titulo: Optional[str] = ""  # Título do artigo (contexto)


# ============================================================================
# HELPERS - GERAIS
# ============================================================================

async def obter_docx_bytes(docx_url: Optional[str], docx_base64: Optional[str], http_client=None) -> bytes:
    """Obtém bytes do DOCX a partir de URL ou base64."""
    if docx_base64:
        import base64
        # Remove prefixo data:... se presente
        if docx_base64.startswith("data:"):
            # Formato: data:application/...;base64,XXXXXX
            docx_base64 = docx_base64.split(",", 1)[1]
        return base64.b64decode(docx_base64)
    elif docx_url:
        if http_client is None:
            async with httpx.AsyncClient(timeout=60) as client:
                resp = await client.get(docx_url)
                resp.raise_for_status()
                return resp.content
        else:
            resp = await http_client.get(docx_url)
            resp.raise_for_status()
            return resp.content
    else:
        raise ValueError("Deve fornecer docx_url ou docx_base64")


def gerar_codigo_cursos(nome_curso: str) -> str:
    nome = unidecode(nome_curso)
    nome = nome.lower()
    nome = re.sub(r'[^a-z0-9 ]', '', nome)
    codigo = re.sub(r'\s+', '-', nome).strip('-')
    return codigo


def rolar_e_coletar_vagas(page, container_locator, max_rolagens=30, pausa=1.0):
    vagas_coletadas = set()
    for _ in range(max_rolagens):
        container_locator.evaluate("el => el.scrollBy(0, 1000)")
        time.sleep(pausa)
        soup = BeautifulSoup(page.content(), "html.parser")
        novos_links = {
            a["href"].split("?")[0]
            for a in soup.select('a[href^="/jobs/view/"]')
            if "href" in a.attrs
        }
        antes = len(vagas_coletadas)
        vagas_coletadas.update(novos_links)
        if len(vagas_coletadas) == antes:
            break
    return list(vagas_coletadas)


def remover_emojis_e_simbolos(texto):
    return ''.join(
        c for c in texto
        if not unicodedata.category(c).startswith("So")
        and not unicodedata.category(c).startswith("Sk")
    )


def remover_caracteres_invisiveis(texto):
    invisiveis = ['\u200b', '\u200c', '\u200d', '\uFEFF']
    for c in invisiveis:
        texto = texto.replace(c, '')
    return texto


def limpar_texto(texto):
    texto = texto.strip()
    texto = re.sub(r"\s+", " ", texto)
    texto = remover_caracteres_invisiveis(texto)
    texto = remover_emojis_e_simbolos(texto)
    return texto


def login_alura(page, user: str, password: str):
    page.goto("https://cursos.alura.com.br/loginForm")
    page.fill("#login-email", user)
    page.fill("#password", password)
    page.click("button:has-text('Entrar')")
    time.sleep(10)
    print("✅ Login realizado com sucesso na Alura.")


def login_linkedin(page, user: str, password: str):
    page.goto("https://www.linkedin.com/checkpoint/lg/sign-in-another-account")
    page.fill("input#username", user)
    page.fill("input#password", password)
    page.click("button[type='submit']")
    time.sleep(10)
    print("✅ Login realizado com sucesso no LinkedIn.")


# ============================================================================
# HELPERS - EXTRAÇÃO DE ARTIGOS (BeautifulSoup)
# ============================================================================

def is_banner_or_promotional(element):
    """Verifica se elemento é banner/propaganda."""
    parent_a = element.find_parent('a') if element.name != 'a' else element
    if parent_a and parent_a.get('href'):
        href = parent_a.get('href', '')
        promo_patterns = [
            '/escola-', '/formacao-', '/planos-', '/curso-online',
            '/empresas', 'cursos.alura.com.br/loginForm',
            'utm_source=blog', 'utm_medium=banner', 'utm_campaign=',
            '/carreiras/', '/pos-tech'
        ]
        for pattern in promo_patterns:
            if pattern in href:
                return True
    
    if element.name == 'img':
        src = element.get('src', '').lower()
        alt = element.get('alt', '').lower()
        if any(x in src for x in ['matricula-escola', 'saiba-mais', 'banner']):
            return True
        if 'banner' in alt:
            return True
    
    return False


def is_site_chrome(element):
    """Verifica se elemento faz parte do chrome do site."""
    if element.find_parent(['nav', 'footer', 'aside']):
        return True
    
    parent_header = element.find_parent('header')
    if parent_header:
        if parent_header.find('a', href=lambda x: x and '/carreiras' in x):
            return True
    
    if element.find_parent(class_=lambda x: x and 'cosmos-author' in str(x)):
        return True
    
    if element.find_parent(class_=lambda x: x and 'social-media' in str(x)):
        return True
    if element.find_parent(class_=lambda x: x and 'cosmos-container-social' in str(x)):
        return True
    
    if element.name == 'p':
        text = element.get_text(strip=True).lower()
        if text == 'compartilhe':
            return True
    
    return False


def is_decorative_element(element):
    """Verifica se é elemento decorativo."""
    if element.name == 'img':
        src = element.get('src', '').lower()
        alt = element.get('alt', '').lower()
        classes = element.get('class', [])
        
        if 'cosmos-image' in classes:
            return False
        
        if 'cdn-wcsm.alura.com.br' in src:
            return False
        
        decorative_patterns = [
            '/assets/img/header/', '/assets/img/home/', '/assets/img/caelum',
            '/assets/img/footer/', '/assets/img/ecossistema/',
            'arrow-', 'return-', 'icon', 'avatar', 
            'gravatar.com/avatar', 'gnarususercontent.com.br'
        ]
        
        for pattern in decorative_patterns:
            if pattern in src:
                return True
        
        if '.svg' in src and '/assets/' in src:
            return True
        
        width = element.get('width')
        if width:
            try:
                if int(width) < 50:
                    return True
            except ValueError:
                pass
    
    return False


def get_text_preserving_spaces(element):
    """Extrai texto preservando espaços entre elementos inline."""
    texts = []
    for child in element.descendants:
        if isinstance(child, NavigableString):
            texts.append(str(child))
    result = ''.join(texts)
    result = re.sub(r'\s+', ' ', result)
    return result.strip()


def extract_text_with_formatting(element, base_url):
    """Extrai texto preservando formatação (links, bold, italic)."""
    segments = []
    
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                segments.append({"text": text})
        
        elif child.name == 'a':
            href = child.get('href', '')
            text = child.get_text()
            if text.strip():
                if href and not href.startswith('http') and not href.startswith('#'):
                    href = urljoin(base_url, href)
                segments.append({"text": text, "link": href if href else None})
        
        elif child.name in ['strong', 'b']:
            inner_a = child.find('a')
            inner_em = child.find(['em', 'i'])
            
            if inner_a:
                href = inner_a.get('href', '')
                if href and not href.startswith('http') and not href.startswith('#'):
                    href = urljoin(base_url, href)
                text = child.get_text()
                if text.strip():
                    segments.append({"text": text, "link": href, "bold": True})
            elif inner_em:
                for subchild in child.children:
                    if isinstance(subchild, NavigableString):
                        text = str(subchild)
                        if text.strip():
                            segments.append({"text": text, "bold": True})
                    elif subchild.name in ['em', 'i']:
                        em_a = subchild.find('a')
                        if em_a:
                            href = em_a.get('href', '')
                            if href and not href.startswith('http') and not href.startswith('#'):
                                href = urljoin(base_url, href)
                            segments.append({"text": subchild.get_text(), "link": href, "bold": True, "italic": True})
                        else:
                            segments.append({"text": subchild.get_text(), "bold": True, "italic": True})
                    elif subchild.name == 'a':
                        href = subchild.get('href', '')
                        if href and not href.startswith('http') and not href.startswith('#'):
                            href = urljoin(base_url, href)
                        segments.append({"text": subchild.get_text(), "link": href, "bold": True})
            else:
                text = child.get_text()
                if text.strip():
                    segments.append({"text": text, "bold": True})
        
        elif child.name in ['em', 'i']:
            inner_a = child.find('a')
            if inner_a:
                href = inner_a.get('href', '')
                if href and not href.startswith('http') and not href.startswith('#'):
                    href = urljoin(base_url, href)
                text = child.get_text()
                if text.strip():
                    segments.append({"text": text, "link": href, "italic": True})
            else:
                text = child.get_text()
                if text.strip():
                    segments.append({"text": text, "italic": True})
        
        elif child.name == 'code':
            text = child.get_text()
            if text.strip():
                segments.append({"text": f"`{text}`", "bold": True})
        
        elif child.name == 'p':
            inner_segments = extract_text_with_formatting(child, base_url)
            segments.extend(inner_segments)
        
        elif child.name in ['span', 'mark', 'u']:
            inner_segments = extract_text_with_formatting(child, base_url)
            segments.extend(inner_segments)
        
        elif child.name == 'br':
            segments.append({"text": "\n"})
        
        elif child.name in ['sup', 'sub']:
            text = child.get_text()
            if text.strip():
                segments.append({"text": text})
        
        else:
            text = child.get_text()
            if text.strip():
                segments.append({"text": text})
    
    return segments


def process_list_items(ul_or_ol, base_url, ordered=False):
    """Processa itens de lista, incluindo listas aninhadas."""
    items = []
    
    for li in ul_or_ol.find_all('li', recursive=False):
        item = {}
        sublist = li.find(['ul', 'ol'], recursive=False)
        
        if sublist:
            sublist_copy = sublist.extract()
            segments = extract_text_with_formatting(li, base_url)
            li.append(sublist_copy)
            
            if segments:
                has_formatting = any(
                    seg.get('link') or seg.get('bold') or seg.get('italic')
                    for seg in segments
                )
                
                if has_formatting:
                    item['segments'] = segments
                elif len(segments) == 1:
                    item['text'] = segments[0].get('text', '').strip()
                else:
                    item['text'] = ''.join(seg.get('text', '') for seg in segments).strip()
            
            sub_ordered = sublist_copy.name == 'ol'
            sub_items = process_list_items(sublist_copy, base_url, sub_ordered)
            if sub_items:
                item['sublist'] = {
                    'ordered': sub_ordered,
                    'items': sub_items
                }
        else:
            segments = extract_text_with_formatting(li, base_url)
            if segments:
                has_formatting = any(
                    seg.get('link') or seg.get('bold') or seg.get('italic')
                    for seg in segments
                )
                
                if has_formatting:
                    item['segments'] = segments
                elif len(segments) == 1:
                    item['text'] = segments[0].get('text', '').strip()
                else:
                    item['text'] = ''.join(seg.get('text', '') for seg in segments).strip()
        
        if item:
            items.append(item)
    
    return items


def extract_table(table_tag):
    """Extrai dados de tabela HTML."""
    headers = []
    rows = []
    
    thead = table_tag.find('thead')
    if thead:
        header_row = thead.find('tr')
        if header_row:
            headers = [th.get_text(strip=True) for th in header_row.find_all(['th', 'td'])]
    
    if not headers:
        first_row = table_tag.find('tr')
        if first_row:
            ths = first_row.find_all('th')
            if ths:
                headers = [th.get_text(strip=True) for th in ths]
    
    tbody = table_tag.find('tbody') or table_tag
    for tr in tbody.find_all('tr'):
        if tr.find('th') and not rows and headers:
            continue
        
        cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
        if cells and any(c for c in cells):
            rows.append(cells)
    
    return headers, rows


def extract_article_content(html: str, base_url: str) -> dict:
    """
    Extrai conteúdo estruturado de artigo Alura usando BeautifulSoup.
    100% determinístico, sem IA!
    """
    soup = BeautifulSoup(html, 'html.parser')
    
    for tag in soup.find_all(['script', 'style', 'noscript', 'svg', 'iframe']):
        tag.decompose()
    
    metadata = {
        'title': None,
        'author': None,
        'publishDate': None
    }
    content = []
    processed_elements = set()
    
    h1 = soup.find('h1')
    if h1:
        metadata['title'] = h1.get_text(strip=True)
        processed_elements.add(id(h1))
    
    date_pattern = re.compile(r'\d{2}/\d{2}/\d{4}')
    page_text = soup.get_text()
    date_match = date_pattern.search(page_text)
    if date_match:
        metadata['publishDate'] = date_match.group()
    
    author_candidates = []
    for img in soup.find_all('img'):
        src = img.get('src', '')
        alt = img.get('alt', '')
        if 'gravatar.com' in src or 'gnarususercontent.com.br' in src:
            if alt and len(alt) > 2 and not any(x in alt.lower() for x in ['logo', 'banner', 'alura']):
                author_candidates.append(alt)
    
    if author_candidates:
        metadata['author'] = author_candidates[0]
    
    main_content = soup.find('body') or soup
    stop_processing = False
    
    list_item_texts = set()
    for li in main_content.find_all('li'):
        li_text = li.get_text(strip=True)
        if li_text and len(li_text) > 10:
            list_item_texts.add(li_text)
    
    for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'p', 'ul', 'ol', 
                                           'blockquote', 'pre', 'table', 'img', 'figure']):
        elem_id = id(element)
        if elem_id in processed_elements:
            continue
        processed_elements.add(elem_id)
        
        if is_site_chrome(element):
            continue
        if is_banner_or_promotional(element):
            continue
        if is_decorative_element(element):
            continue
        
        if element.name in ['h2', 'h3']:
            text = element.get_text(strip=True).lower()
            if any(x in text for x in ['leia também', 'artigos relacionados', 'veja outros artigos']):
                stop_processing = True
        
        if stop_processing:
            continue
        
        if element.name == 'h1':
            continue
        
        if element.name in ['h2', 'h3', 'h4', 'h5']:
            text = get_text_preserving_spaces(element)
            if text and len(text) > 1:
                if element.find_parent(class_=lambda x: x and 'toc' in x.lower() if x else False):
                    continue
                level = int(element.name[1])
                content.append({
                    'type': 'heading',
                    'level': level,
                    'text': text
                })
        
        elif element.name == 'p':
            text = element.get_text(strip=True)
            if not text:
                continue
            if text in list_item_texts:
                continue
            
            segments = extract_text_with_formatting(element, base_url)
            if segments:
                has_formatting = any(
                    seg.get('link') or seg.get('bold') or seg.get('italic') 
                    for seg in segments
                )
                
                if not has_formatting and len(segments) == 1:
                    content.append({
                        'type': 'paragraph',
                        'text': segments[0].get('text', '').strip()
                    })
                else:
                    content.append({
                        'type': 'paragraph',
                        'segments': segments
                    })
        
        elif element.name in ['ul', 'ol']:
            if element.find_parent(['ul', 'ol']):
                continue
            
            ordered = element.name == 'ol'
            items = process_list_items(element, base_url, ordered)
            
            if items:
                content.append({
                    'type': 'list',
                    'ordered': ordered,
                    'items': items
                })
        
        elif element.name == 'blockquote':
            segments = extract_text_with_formatting(element, base_url)
            cite_tag = element.find('cite')
            cite = cite_tag.get_text(strip=True) if cite_tag else None
            
            if segments:
                blockquote_item = {'type': 'blockquote', 'segments': segments}
                if cite:
                    blockquote_item['cite'] = cite
                content.append(blockquote_item)
        
        elif element.name == 'pre':
            code_tag = element.find('code')
            if code_tag:
                code_content = code_tag.get_text()
                classes = code_tag.get('class', [])
                language = None
                for cls in classes:
                    if isinstance(cls, str):
                        if cls.startswith('language-'):
                            language = cls.replace('language-', '')
                            break
                        elif cls in ['python', 'javascript', 'java', 'sql', 'bash', 
                                    'html', 'css', 'json', 'typescript', 'jsx', 'ruby',
                                    'go', 'rust', 'php', 'csharp', 'kotlin', 'swift']:
                            language = cls
                            break
                
                content.append({
                    'type': 'code',
                    'language': language,
                    'content': code_content
                })
            else:
                content.append({
                    'type': 'code',
                    'content': element.get_text()
                })
        
        elif element.name == 'table':
            headers, rows = extract_table(element)
            if headers or rows:
                content.append({
                    'type': 'table',
                    'headers': headers,
                    'rows': rows
                })
        
        elif element.name == 'img':
            src = element.get('src', '')
            if not src:
                continue
            if is_banner_or_promotional(element):
                continue
            if is_decorative_element(element):
                continue
            
            if not src.startswith('http'):
                src = urljoin(base_url, src)
            
            alt = element.get('alt', '')
            width = element.get('width')
            height = element.get('height')
            
            img_item = {
                'type': 'image',
                'url': src,
                'alt': alt
            }
            
            if width:
                try:
                    img_item['width'] = int(width)
                except:
                    pass
            if height:
                try:
                    img_item['height'] = int(height)
                except:
                    pass
            
            content.append(img_item)
        
        elif element.name == 'figure':
            img = element.find('img')
            if img:
                src = img.get('src', '')
                if not src:
                    continue
                
                if not src.startswith('http'):
                    src = urljoin(base_url, src)
                
                figcaption = element.find('figcaption')
                alt = figcaption.get_text(strip=True) if figcaption else img.get('alt', '')
                
                content.append({
                    'type': 'image',
                    'url': src,
                    'alt': alt
                })
                processed_elements.add(id(img))
    
    content = [item for item in content if item]
    
    stats = {}
    for item in content:
        item_type = item.get('type', 'unknown')
        stats[item_type] = stats.get(item_type, 0) + 1
    
    filename = metadata.get('title', 'documento') or 'documento'
    filename = unidecode(filename)
    filename = re.sub(r'[^a-zA-Z0-9\s-]', '', filename)
    filename = re.sub(r'\s+', '-', filename).strip('-')
    filename = filename[:80]
    filename = f"{filename}.docx"
    
    return {
        'metadata': metadata,
        'content': content,
        'filename': filename,
        'base_url': base_url,
        'stats': stats
    }


# ============================================================================
# HELPERS - GERAÇÃO DE DOCX
# ============================================================================

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0066CC')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def convert_relative_url(url: str, base_url: str) -> str:
    if not url:
        return url
    if url.startswith('http://') or url.startswith('https://'):
        return url
    if not base_url:
        return url
    return urljoin(base_url, url)


def download_image(url: str) -> Optional[BytesIO]:
    try:
        with httpx.Client(timeout=30, follow_redirects=True) as client:
            response = client.get(url)
            response.raise_for_status()
            return BytesIO(response.content)
    except Exception as e:
        print(f"❌ Erro ao baixar imagem {url}: {e}")
        return None


def convert_image_for_docx(image_bytes: Optional[BytesIO]) -> Optional[BytesIO]:
    """
    Converte imagem para formato compativel com python-docx.

    - WEBP animado -> GIF (preserva animacao)
    - WEBP estatico -> PNG
    - Outros formatos nao suportados -> PNG
    - Formatos suportados (PNG, JPEG, GIF, BMP, TIFF) -> retorna original
    """
    if image_bytes is None:
        return None

    SUPPORTED_FORMATS = {'PNG', 'JPEG', 'GIF', 'BMP', 'TIFF', 'JPG'}

    try:
        image_bytes.seek(0)
        img = PILImage.open(image_bytes)

        # Se ja e formato suportado, retorna original
        if img.format and img.format.upper() in SUPPORTED_FORMATS:
            image_bytes.seek(0)
            return image_bytes

        # WEBP animado -> GIF
        if img.format == 'WEBP' and getattr(img, 'is_animated', False):
            print(f"  [CONV] WEBP animado -> GIF ({img.n_frames} frames)")
            return _convert_animated_webp_to_gif(img)

        # WEBP estatico ou outro formato -> PNG
        print(f"  [CONV] {img.format or 'unknown'} -> PNG")
        return _convert_to_png(img)

    except Exception as e:
        print(f"  [ERRO] Conversao de imagem: {e}")
        return None


def _convert_animated_webp_to_gif(img) -> BytesIO:
    """Converte WEBP animado para GIF preservando animacao."""
    frames = []
    durations = []

    try:
        while True:
            frame = img.copy()

            # GIF requer modo P (paleta) ou L (grayscale)
            if frame.mode in ('RGBA', 'LA'):
                # Compoe sobre fundo branco para remover transparencia
                background = PILImage.new('RGBA', frame.size, (255, 255, 255, 255))
                if frame.mode == 'RGBA':
                    background.paste(frame, mask=frame.split()[3])
                else:
                    background.paste(frame)
                frame = background.convert('RGB').convert('P', palette=PILImage.ADAPTIVE, colors=256)
            elif frame.mode != 'P':
                frame = frame.convert('P', palette=PILImage.ADAPTIVE, colors=256)

            frames.append(frame)
            durations.append(img.info.get('duration', 100))
            img.seek(img.tell() + 1)
    except EOFError:
        pass

    if not frames:
        raise ValueError("Nenhum frame extraido do WEBP animado")

    gif_buffer = BytesIO()
    frames[0].save(
        gif_buffer,
        format='GIF',
        save_all=True,
        append_images=frames[1:],
        duration=durations,
        loop=0
    )
    gif_buffer.seek(0)
    return gif_buffer


def _convert_to_png(img) -> BytesIO:
    """Converte imagem para PNG."""
    # Converte para RGB se necessario (remove canal alpha para compatibilidade)
    if img.mode in ('RGBA', 'LA', 'P'):
        # Mantem RGBA para PNG (suporta transparencia)
        if img.mode == 'P':
            img = img.convert('RGBA')
    elif img.mode not in ('RGB', 'L'):
        img = img.convert('RGB')

    png_buffer = BytesIO()
    img.save(png_buffer, format='PNG')
    png_buffer.seek(0)
    return png_buffer


def get_image_dimensions_from_bytes(image_bytes: BytesIO) -> tuple:
    try:
        image_bytes.seek(0)
        img = PILImage.open(image_bytes)
        width, height = img.size
        image_bytes.seek(0)
        return width, height
    except:
        return None, None


def set_paragraph_shading(paragraph, color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    paragraph._p.get_or_add_pPr().append(shading)


def add_left_border(paragraph, color: str = '0066CC', width: int = 24):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(width))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def process_list_item_content_docx(doc, li, paragraph):
    """Processa conteúdo de item de lista no DOCX."""
    if li is None:
        return
    if isinstance(li, dict):
        if 'segments' in li and li['segments']:
            for seg in li['segments']:
                if seg is None or not isinstance(seg, dict):
                    continue
                seg_text = seg.get('text', '') or ''
                seg_link = seg.get('link')
                seg_bold = seg.get('bold', False)
                seg_italic = seg.get('italic', False)

                if seg_link:
                    add_hyperlink(paragraph, seg_text, seg_link)
                else:
                    run = paragraph.add_run(seg_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    if seg_bold:
                        run.bold = True
                    if seg_italic:
                        run.italic = True
        elif 'text' in li and li['text']:
            run = paragraph.add_run(str(li['text']))
            run.font.name = 'Arial'
            run.font.size = Pt(12)
    elif li:
        run = paragraph.add_run(str(li))
        run.font.name = 'Arial'
        run.font.size = Pt(12)


def process_nested_list_docx(doc, items, ordered=False, indent_level=0):
    """Processa lista aninhada no DOCX."""
    if not items:
        return
    markers = ["• ", "◦ ", "▪ ", "- "]

    for idx, li in enumerate(items):
        if li is None:
            continue
        list_para = doc.add_paragraph()
        list_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if ordered:
            prefix = f"{idx + 1}. "
        else:
            prefix = markers[min(indent_level, len(markers) - 1)]
        
        prefix_run = list_para.add_run(prefix)
        prefix_run.font.name = 'Arial'
        prefix_run.font.size = Pt(12)
        
        process_list_item_content_docx(doc, li, list_para)
        
        base_indent = 0.5
        list_para.paragraph_format.left_indent = Inches(base_indent + (indent_level * 0.3))
        list_para.space_after = Pt(3)
        
        if isinstance(li, dict) and 'sublist' in li and li['sublist']:
            sublist = li['sublist']
            sub_ordered = sublist.get('ordered', False)
            sub_items = sublist.get('items', [])
            if sub_items:
                process_nested_list_docx(doc, sub_items, sub_ordered, indent_level + 1)


# ============================================================================
# HELPERS - PROCESSAMENTO DE VÍDEO
# ============================================================================

def criar_video_com_transicoes(videos, audio_narracao, output, transicao_duracao=0.5, transicao_tipo="fade", legendas_srt=None, estilo_legenda="youtube", legenda_config=None):
    if len(videos) == 0:
        raise ValueError("Nenhum vídeo fornecido")
    
    temp_video_sem_audio = output.replace('.mp4', '_temp.mp4')
    
    try:
        if len(videos) == 1:
            shutil.copy(videos[0], temp_video_sem_audio)
        else:
            print(f"🔄 Juntando {len(videos)} vídeos com transições...")
            filter_parts = []
            last_label = "[0:v]"
            for i in range(len(videos) - 1):
                next_input = f"[{i+1}:v]"
                out_label = f"[v{i}]" if i < len(videos) - 2 else "[vout]"
                offset = (i + 1) * 5 - transicao_duracao
                xfade = f"{last_label}{next_input}xfade=transition={transicao_tipo}:duration={transicao_duracao}:offset={offset}{out_label}"
                filter_parts.append(xfade)
                last_label = out_label
            filter_complex = ";".join(filter_parts)
            cmd = ['ffmpeg', '-y']
            for video in videos:
                cmd.extend(['-i', video])
            cmd.extend(['-filter_complex', filter_complex, '-map', '[vout]', '-c:v', 'libx264', '-preset', 'faster', '-pix_fmt', 'yuv420p', '-an', temp_video_sem_audio])
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode != 0:
                raise Exception(f"Erro ao juntar vídeos: {result.stderr}")
        
        print(f"🔄 Adicionando áudio da narração...")
        
        def get_duration(file_path):
            cmd = ['ffprobe', '-v', 'error', '-show_entries', 'format=duration', '-of', 'default=noprint_wrappers=1:nokey=1', file_path]
            result = subprocess.run(cmd, capture_output=True, text=True)
            return float(result.stdout.strip())
        
        video_duration = get_duration(temp_video_sem_audio)
        audio_duration = get_duration(audio_narracao)
        
        estilos_predefinidos = {
            "youtube": "FontName=Arial Black,FontSize=28,Bold=1,PrimaryColour=&HFFFFFF,OutlineColour=&H000000,BackColour=&H80000000,Outline=3,Shadow=2,MarginV=40",
            "discreto": "FontName=Arial,FontSize=18,PrimaryColour=&HFFFFFF,OutlineColour=&H000000,Outline=1,MarginV=20"
        }
        
        if estilo_legenda == "custom" and legenda_config:
            style = f"FontName={legenda_config.font_name},FontSize={legenda_config.font_size},Bold={1 if legenda_config.bold else 0},PrimaryColour={legenda_config.primary_colour},OutlineColour={legenda_config.outline_colour},BackColour={legenda_config.back_colour},Outline={legenda_config.outline},Shadow={legenda_config.shadow},MarginV={legenda_config.margin_v}"
        else:
            style = estilos_predefinidos.get(estilo_legenda, estilos_predefinidos["youtube"])
        
        if audio_duration > video_duration:
            diff = audio_duration - video_duration
            fade_duration = min(1.0, diff)
            fade_start = video_duration - fade_duration
            if legendas_srt:
                srt_escaped = legendas_srt.replace('\\', '/').replace(':', '\\:')
                filter_complex = f"[0:v]fade=t=out:st={fade_start}:d={fade_duration},tpad=stop_mode=add:stop_duration={diff}:color=black,subtitles={srt_escaped}:force_style='{style}'[v]"
            else:
                filter_complex = f'[0:v]fade=t=out:st={fade_start}:d={fade_duration},tpad=stop_mode=add:stop_duration={diff}:color=black[v]'
            cmd = ['ffmpeg', '-y', '-i', temp_video_sem_audio, '-i', audio_narracao, '-filter_complex', filter_complex, '-map', '[v]', '-map', '1:a:0', '-c:v', 'libx264', '-preset', 'faster', '-c:a', 'aac', '-b:a', '192k', '-pix_fmt', 'yuv420p', output]
        else:
            if legendas_srt:
                srt_escaped = legendas_srt.replace('\\', '/').replace(':', '\\:')
                cmd = ['ffmpeg', '-y', '-i', temp_video_sem_audio, '-i', audio_narracao, '-vf', f"subtitles={srt_escaped}:force_style='{style}'", '-c:v', 'libx264', '-preset', 'faster', '-c:a', 'aac', '-b:a', '192k', '-pix_fmt', 'yuv420p', '-shortest', output]
            else:
                cmd = ['ffmpeg', '-y', '-i', temp_video_sem_audio, '-i', audio_narracao, '-c:v', 'copy', '-c:a', 'aac', '-b:a', '192k', '-shortest', output]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise Exception(f"Erro ao adicionar áudio: {result.stderr}")
        
        print(f"✅ Vídeo processado!")
    finally:
        if os.path.exists(temp_video_sem_audio):
            os.remove(temp_video_sem_audio)


def gerar_legendas_srt(audio_path, output_srt):
    print(f"🎙️ Transcrevendo áudio com Whisper...")
    try:
        with open(audio_path, "rb") as audio_file:
            transcript = client.audio.transcriptions.create(model="whisper-1", file=audio_file, response_format="srt", language="pt")
        with open(output_srt, "w", encoding="utf-8") as f:
            f.write(transcript)
        print(f"✅ Legendas geradas: {output_srt}")
        return output_srt
    except Exception as e:
        raise Exception(f"Erro ao transcrever áudio: {str(e)}")


def cleanup_job(job_dir, delay_seconds=3600):
    time.sleep(delay_seconds)
    if job_dir.exists():
        shutil.rmtree(job_dir, ignore_errors=True)


def baixar_arquivo(url, destino):
    response = requests.get(url, stream=True, timeout=60)
    response.raise_for_status()
    with open(destino, 'wb') as f:
        for chunk in response.iter_content(chunk_size=8192):
            f.write(chunk)


# ============================================================================
# HELPERS - LIBREOFFICE
# ============================================================================

class LibreOfficeConnection:
    """Singleton para conexão com LibreOffice."""
    
    _desktop = None
    
    @classmethod
    def get_desktop(cls):
        if cls._desktop is None:
            cls._connect()
        return cls._desktop
    
    @classmethod
    def _connect(cls, host="127.0.0.1", port=2002):
        if not LIBREOFFICE_DISPONIVEL:
            raise RuntimeError("python3-uno não instalado")
        
        local_context = uno.getComponentContext()
        resolver = local_context.ServiceManager.createInstanceWithContext(
            "com.sun.star.bridge.UnoUrlResolver", local_context
        )
        
        for attempt in range(3):
            try:
                ctx = resolver.resolve(
                    f"uno:socket,host={host},port={port};urp;StarOffice.ComponentContext"
                )
                smgr = ctx.ServiceManager
                cls._desktop = smgr.createInstanceWithContext(
                    "com.sun.star.frame.Desktop", ctx
                )
                return
            except Exception as e:
                if attempt < 2:
                    time.sleep(2)
                else:
                    raise RuntimeError(f"Não conectou ao LibreOffice: {e}")
    
    @classmethod
    def reset(cls):
        """Reset da conexão (usar se LibreOffice reiniciar)."""
        cls._desktop = None


def _extrair_texto_lo(docx_path: str) -> dict:
    """Extrai texto do documento com posições."""
    desktop = LibreOfficeConnection.get_desktop()
    
    url = f"file://{os.path.abspath(docx_path)}"
    props = (PropertyValue("Hidden", 0, True, 0),)
    
    doc = desktop.loadComponentFromURL(url, "_blank", 0, props)
    if not doc:
        raise RuntimeError(f"Não abriu: {docx_path}")
    
    try:
        text = doc.getText()
        enum = text.createEnumeration()
        
        paragrafos = []
        texto_parts = []
        idx = 0
        
        while enum.hasMoreElements():
            element = enum.nextElement()
            if element.supportsService("com.sun.star.text.Paragraph"):
                texto = element.getString()
                paragrafos.append({
                    "indice": idx,
                    "texto": texto,
                    "tamanho": len(texto)
                })
                texto_parts.append(f"[P{idx}] {texto}")
                idx += 1
        
        return {
            "paragrafos": paragrafos,
            "texto_completo": "\n".join(texto_parts),
            "total_paragrafos": idx
        }
    finally:
        doc.close(True)


def _aplicar_revisoes_lo(docx_path: str, revisoes: list, autor: str, output_path: str) -> dict:
    """Aplica revisões usando LibreOffice com busca inteligente e validação."""
    desktop = LibreOfficeConnection.get_desktop()
    
    url = f"file://{os.path.abspath(docx_path)}"
    props = (PropertyValue("Hidden", 0, True, 0),)
    
    doc = desktop.loadComponentFromURL(url, "_blank", 0, props)
    if not doc:
        raise RuntimeError(f"Não abriu: {docx_path}")
    
    try:
        # Ativa Track Changes
        doc.setPropertyValue("RecordChanges", True)
        doc.setPropertyValue("ShowChanges", True)
        
        text = doc.getText()
        enum = text.createEnumeration()
        
        # Extrai parágrafos E seus textos
        paragrafos = []
        textos_paragrafos = []
        while enum.hasMoreElements():
            element = enum.nextElement()
            if element.supportsService("com.sun.star.text.Paragraph"):
                texto = element.getString()
                paragrafos.append(element)
                textos_paragrafos.append(texto)
        
        # Aplica revisões em ordem reversa
        revisoes_ord = sorted(
            enumerate(revisoes),
            key=lambda x: (x[1].paragrafo, x[1].inicio),
            reverse=True
        )
        
        resultados = []
        
        for idx_orig, rev in revisoes_ord:
            try:
                if rev.paragrafo >= len(paragrafos):
                    resultados.append({
                        "idx": idx_orig,
                        "ok": False,
                        "erro": f"Parágrafo {rev.paragrafo} não existe (máx: {len(paragrafos)-1})"
                    })
                    continue
                
                para = paragrafos[rev.paragrafo]
                texto_para = textos_paragrafos[rev.paragrafo]
                
                # VALIDAÇÃO E BUSCA INTELIGENTE
                inicio_real = rev.inicio
                fim_real = rev.fim
                
                # Verifica se as posições são válidas
                if inicio_real < 0 or fim_real > len(texto_para) or inicio_real > fim_real:
                    resultados.append({
                        "idx": idx_orig,
                        "ok": False,
                        "erro": f"Posições inválidas: [{inicio_real}:{fim_real}] (tamanho: {len(texto_para)})"
                    })
                    continue
                
                # Se tem texto_esperado, valida e procura se necessário
                if hasattr(rev, 'texto_esperado') and rev.texto_esperado and rev.acao in ["substituir", "deletar", "comentario"]:
                    texto_atual = texto_para[inicio_real:fim_real]
                    
                    # Compara ignorando espaços extras
                    texto_atual_norm = ' '.join(texto_atual.split())
                    texto_esperado_norm = ' '.join(rev.texto_esperado.split())
                    
                    if texto_atual_norm != texto_esperado_norm:
                        # Texto não corresponde! Tenta encontrar
                        print(f"⚠️ Revisão {idx_orig}: texto não corresponde. Buscando...")
                        print(f"  Esperado: '{texto_esperado_norm[:50]}...'")
                        print(f"  Atual: '{texto_atual_norm[:50]}...'")
                        
                        # Procura o texto esperado no parágrafo
                        pos_encontrada = texto_para.find(rev.texto_esperado)
                        
                        if pos_encontrada == -1:
                            # Tenta busca fuzzy (ignorando espaços)
                            texto_para_norm = ' '.join(texto_para.split())
                            pos_encontrada_norm = texto_para_norm.find(texto_esperado_norm)
                            
                            if pos_encontrada_norm != -1:
                                # Encontrou! Ajusta posição aproximada
                                # (não é perfeito mas é melhor que nada)
                                inicio_real = max(0, pos_encontrada_norm - 10)
                                fim_real = min(len(texto_para), pos_encontrada_norm + len(texto_esperado_norm) + 10)
                                print(f"✅ Encontrado (fuzzy): aproximado em [{inicio_real}:{fim_real}]")
                            else:
                                resultados.append({
                                    "idx": idx_orig,
                                    "ok": False,
                                    "erro": f"Texto esperado não encontrado: '{rev.texto_esperado[:30]}...'"
                                })
                                continue
                        else:
                            inicio_real = pos_encontrada
                            fim_real = pos_encontrada + len(rev.texto_esperado)
                            print(f"✅ Encontrado: [{inicio_real}:{fim_real}]")
                
                # Cria cursor na posição correta
                cursor = para.getText().createTextCursor()
                cursor.gotoStart(False)
                
                # Move cursor para a posição
                if inicio_real > 0:
                    cursor.goRight(inicio_real, False)
                
                # Aplica a ação
                if rev.acao == "substituir":
                    tamanho_selecao = fim_real - inicio_real
                    if tamanho_selecao > 0:
                        cursor.goRight(tamanho_selecao, True)
                    cursor.setString(rev.texto_novo)
                    _add_comment_lo(doc, cursor, f"[{rev.tipo}] {rev.justificativa}", autor)
                    resultados.append({"idx": idx_orig, "ok": True})
                
                elif rev.acao == "deletar":
                    tamanho_selecao = fim_real - inicio_real
                    if tamanho_selecao > 0:
                        cursor.goRight(tamanho_selecao, True)
                    cursor.setString("")
                    _add_comment_lo(doc, cursor, f"[{rev.tipo}] Removido: {rev.justificativa}", autor)
                    resultados.append({"idx": idx_orig, "ok": True})
                
                elif rev.acao == "inserir":
                    cursor.getText().insertString(cursor, rev.texto_novo, False)
                    _add_comment_lo(doc, cursor, f"[{rev.tipo}] Inserido: {rev.justificativa}", autor)
                    resultados.append({"idx": idx_orig, "ok": True})
                
                elif rev.acao == "comentario":
                    tamanho_selecao = fim_real - inicio_real
                    if tamanho_selecao > 0:
                        cursor.goRight(tamanho_selecao, True)
                    _add_comment_lo(doc, cursor, f"[{rev.tipo}] {rev.justificativa}", autor)
                    resultados.append({"idx": idx_orig, "ok": True})
                
                else:
                    resultados.append({
                        "idx": idx_orig,
                        "ok": False,
                        "erro": f"Ação desconhecida: {rev.acao}"
                    })
            
            except Exception as e:
                import traceback
                traceback.print_exc()
                resultados.append({
                    "idx": idx_orig,
                    "ok": False,
                    "erro": f"Exceção: {str(e)}"
                })
        
        # Salva documento
        output_url = f"file://{os.path.abspath(output_path)}"
        save_props = (
            PropertyValue("FilterName", 0, "MS Word 2007 XML", 0),
            PropertyValue("Overwrite", 0, True, 0),
        )
        doc.storeToURL(output_url, save_props)
        
        return {
            "arquivo": output_path,
            "total": len(revisoes),
            "ok": sum(1 for r in resultados if r.get("ok")),
            "falhas": sum(1 for r in resultados if not r.get("ok")),
            "detalhes": sorted(resultados, key=lambda x: x["idx"])
        }
    finally:
        doc.close(True)



def _add_comment_lo(doc, cursor, texto: str, autor: str):
    """Adiciona comentário."""
    try:
        ann = doc.createInstance("com.sun.star.text.TextField.Annotation")
        ann.Author = autor
        ann.Content = texto
        cursor.getText().insertTextContent(cursor, ann, False)
    except:
        pass


# ============================================================================
# FASTAPI APP
# ============================================================================

app = FastAPI()


# ============================================================================
# ENDPOINTS - GERAL
# ============================================================================

@app.get("/ping")
def ping():
    return {"ok": True, "service": "runner"}


# ============================================================================
# ENDPOINTS - EXTRAÇÃO E DOCX
# ============================================================================

@app.post("/extract-article")
async def extract_article(payload: ExtractArticlePayload):
    """Extrai conteúdo estruturado de artigo Alura (100% determinístico, SEM IA)."""
    try:
        print(f"📥 Extraindo artigo: {payload.url}")
        
        with httpx.Client(timeout=30, follow_redirects=True) as client:
            response = client.get(payload.url)
            response.raise_for_status()
            html = response.text
        
        print(f"📄 HTML recebido: {len(html)} bytes")
        result = extract_article_content(html, payload.url)
        
        print(f"✅ Extração concluída!")
        print(f"📊 Estatísticas: {result['stats']}")
        
        return result
    
    except httpx.HTTPError as e:
        raise HTTPException(status_code=400, detail=f"Erro ao buscar URL: {str(e)}")
    except Exception as e:
        print(f"❌ Erro na extração: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erro ao extrair artigo: {str(e)}")


@app.post("/html-to-docx")
async def html_to_docx(payload: ExtractArticlePayload):
    """Pipeline completo: extrai artigo e gera DOCX em uma única chamada."""
    try:
        print(f"🚀 Pipeline HTML → DOCX: {payload.url}")
        
        with httpx.Client(timeout=30, follow_redirects=True) as client:
            response = client.get(payload.url)
            response.raise_for_status()
            html = response.text
        
        article_data = extract_article_content(html, payload.url)
        print(f"📊 Extraído: {article_data['stats']}")
        
        docx_payload = GenerateDocxPayload(
            metadata=ArticleMetadata(**article_data['metadata']),
            content=[ContentItem(**item) for item in article_data['content']],
            filename=article_data['filename'],
            base_url=article_data['base_url']
        )
        
        return await generate_docx(docx_payload)
    
    except Exception as e:
        print(f"❌ Erro no pipeline: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erro no pipeline: {str(e)}")


@app.post("/generate-docx")
async def generate_docx(payload: GenerateDocxPayload):
    """Gera documento Word (.docx) a partir de JSON estruturado."""
    try:
        print(f"📝 Gerando DOCX: {payload.metadata.title or 'Sem título'}")
        
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        
        if payload.metadata.title:
            title_para = doc.add_heading(payload.metadata.title, level=1)
            for run in title_para.runs:
                run.bold = True
                run.font.size = Pt(28)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(33, 37, 41)
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.space_after = Pt(6)
        
        meta_parts = []
        if payload.metadata.author:
            meta_parts.append(f"Por {payload.metadata.author}")
        if payload.metadata.publishDate:
            meta_parts.append(payload.metadata.publishDate)
        
        if meta_parts:
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run(" • ".join(meta_parts))
            meta_run.italic = True
            meta_run.font.size = Pt(11)
            meta_run.font.color.rgb = RGBColor(102, 102, 102)
            meta_para.space_after = Pt(12)
        
        doc.add_paragraph("_" * 80)
        
        for item in payload.content:
            if item is None:
                continue

            if item.type == "heading" and item.text:
                spacer = doc.add_paragraph()
                spacer.space_after = Pt(0)
                spacer.space_before = Pt(6)
                
                level = item.level if item.level else 2
                heading_para = doc.add_heading(item.text, level=level)
                
                for run in heading_para.runs:
                    run.bold = True
                    run.font.name = 'Arial'
                    
                    if level == 2:
                        run.font.size = Pt(16)
                        run.font.color.rgb = RGBColor(44, 62, 80)
                    elif level == 3:
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(52, 73, 94)
                    elif level == 4:
                        run.font.size = Pt(13)
                        run.font.color.rgb = RGBColor(60, 80, 100)
                    else:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(70, 90, 110)
                
                heading_para.space_before = Pt(12)
                heading_para.space_after = Pt(6)
            
            elif item.type == "paragraph":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                if item.segments:
                    for seg in item.segments:
                        if seg is None:
                            continue
                        if seg.link:
                            add_hyperlink(para, seg.text or '', seg.link)
                        else:
                            run = para.add_run(seg.text or '')
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            if seg.bold:
                                run.bold = True
                            if seg.italic:
                                run.italic = True
                elif item.text:
                    run = para.add_run(item.text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                
                para.space_after = Pt(6)
            
            elif item.type == "list" and item.items:
                process_nested_list_docx(doc, item.items, item.ordered or False, indent_level=0)
                doc.add_paragraph()
            
            elif item.type == "blockquote":
                if item.segments:
                    quote_para = doc.add_paragraph()
                    quote_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    for seg in item.segments:
                        if seg is None:
                            continue
                        if seg.link:
                            add_hyperlink(quote_para, seg.text or '', seg.link)
                        else:
                            run = quote_para.add_run(seg.text or '')
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            run.italic = True
                            run.font.color.rgb = RGBColor(85, 85, 85)
                    
                    add_left_border(quote_para, color='0066CC', width=24)
                    quote_para.paragraph_format.left_indent = Inches(0.3)
                    quote_para.space_before = Pt(6)
                    quote_para.space_after = Pt(6)
                    
                elif item.text:
                    quote_para = doc.add_paragraph()
                    quote_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    run = quote_para.add_run(item.text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.italic = True
                    run.font.color.rgb = RGBColor(85, 85, 85)
                    
                    add_left_border(quote_para, color='0066CC', width=24)
                    quote_para.paragraph_format.left_indent = Inches(0.3)
                    quote_para.space_before = Pt(6)
                    quote_para.space_after = Pt(6)
                
                if item.cite:
                    cite_para = doc.add_paragraph()
                    cite_run = cite_para.add_run(f"— {item.cite}")
                    cite_run.font.name = 'Arial'
                    cite_run.font.size = Pt(10)
                    cite_run.italic = True
                    cite_run.font.color.rgb = RGBColor(120, 120, 120)
                    cite_para.paragraph_format.left_indent = Inches(0.5)
                    cite_para.space_after = Pt(12)
            
            elif item.type == "code" and item.content:
                if item.language:
                    lang_para = doc.add_paragraph()
                    lang_run = lang_para.add_run(f" {item.language.upper()} ")
                    lang_run.font.name = 'Consolas'
                    lang_run.font.size = Pt(9)
                    lang_run.font.color.rgb = RGBColor(255, 255, 255)
                    set_paragraph_shading(lang_para, '2d2d2d')
                    lang_para.space_after = Pt(0)
                
                for line in item.content.split('\n'):
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run(line if line else ' ')
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(10)
                    code_run.font.color.rgb = RGBColor(51, 51, 51)
                    set_paragraph_shading(code_para, 'F8F8F8')
                    code_para.paragraph_format.left_indent = Inches(0.2)
                    code_para.space_after = Pt(0)
                    code_para.space_before = Pt(0)
                
                doc.add_paragraph().space_after = Pt(12)
            
            elif item.type == "image" and item.url:
                image_url = convert_relative_url(item.url, payload.base_url)
                print(f"🖼️ Baixando imagem: {image_url[:80]}...")
                image_data = download_image(image_url)

                if image_data:
                    # Converte para formato compativel com python-docx se necessario
                    image_data = convert_image_for_docx(image_data)

                if image_data:
                    try:
                        orig_width, orig_height = get_image_dimensions_from_bytes(image_data)
                        max_width_cm = 15

                        if orig_width and orig_height:
                            width_cm = orig_width / 96 * 2.54
                            height_cm = orig_height / 96 * 2.54
                            if width_cm > max_width_cm:
                                ratio = max_width_cm / width_cm
                                width_cm = max_width_cm
                                height_cm = height_cm * ratio
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_para.add_run()
                            run.add_picture(image_data, width=Cm(width_cm))
                        else:
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_para.add_run()
                            run.add_picture(image_data, width=Cm(max_width_cm))

                        img_para.space_after = Pt(6)

                        if item.alt and len(item.alt) > 5:
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run(item.alt)
                            caption_run.italic = True
                            caption_run.font.size = Pt(10)
                            caption_run.font.color.rgb = RGBColor(102, 102, 102)
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.space_after = Pt(12)

                        print(f"✅ Imagem adicionada")
                    except Exception as img_error:
                        print(f"❌ Erro ao processar imagem: {img_error}")
            
            elif item.type == "table" and item.headers and item.rows:
                print(f"📊 Adicionando tabela com {len(item.rows)} linhas...")
                
                num_cols = len(item.headers)
                num_rows = len(item.rows) + 1
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                
                header_row = table.rows[0]
                for idx, header_text in enumerate(item.headers):
                    cell = header_row.cells[idx]
                    cell.text = header_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'E0E0E0')
                    cell._tc.get_or_add_tcPr().append(shading)
                
                for row_idx, row_data in enumerate(item.rows):
                    row = table.rows[row_idx + 1]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = row.cells[col_idx]
                            cell.text = str(cell_text) if cell_text else ""
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                
                doc.add_paragraph().space_after = Pt(12)
        
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        filename = payload.filename
        if not filename.endswith('.docx'):
            filename += '.docx'
        filename = re.sub(r'[^a-zA-Z0-9\s\-_.]', '', filename)
        
        print(f"✅ DOCX gerado: {filename}")
        
        return Response(
            content=doc_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    
    except Exception as e:
        print(f"❌ Erro ao gerar DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erro ao gerar DOCX: {str(e)}")


# ============================================================================
# ENDPOINTS - LIBREOFFICE
# ============================================================================

@app.get("/libreoffice/status")
async def libreoffice_status():
    """Verifica se LibreOffice está disponível."""
    if not LIBREOFFICE_DISPONIVEL:
        return {"status": "indisponivel", "msg": "python3-uno não instalado"}
    
    try:
        LibreOfficeConnection.get_desktop()
        return {"status": "ok", "msg": "LibreOffice conectado"}
    except Exception as e:
        return {"status": "erro", "msg": str(e)}


@app.post("/libreoffice/extrair-texto")
async def libreoffice_extrair_texto(arquivo: UploadFile = File(...)):
    """Extrai texto do documento com posições."""
    if not LIBREOFFICE_DISPONIVEL:
        raise HTTPException(500, "LibreOffice não disponível")
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        content = await arquivo.read()
        tmp.write(content)
        tmp_path = tmp.name
    
    try:
        return _extrair_texto_lo(tmp_path)
    finally:
        os.unlink(tmp_path)


@app.post("/libreoffice/extrair-texto-url")
async def libreoffice_extrair_texto_url(url: str = Form(...)):
    """Extrai texto de documento via URL."""
    if not LIBREOFFICE_DISPONIVEL:
        raise HTTPException(500, "LibreOffice não disponível")
    
    async with httpx.AsyncClient() as client:
        resp = await client.get(url)
        if resp.status_code != 200:
            raise HTTPException(400, f"Erro ao baixar: {resp.status_code}")
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(resp.content)
        tmp_path = tmp.name
    
    try:
        return _extrair_texto_lo(tmp_path)
    finally:
        os.unlink(tmp_path)


@app.post("/libreoffice/aplicar-revisoes")
async def libreoffice_aplicar_revisoes(
    arquivo: UploadFile = File(...),
    revisoes: str = Form(...),
    autor: str = Form("Revisor IA")
):
    """Aplica revisões ao documento com Track Changes real."""
    if not LIBREOFFICE_DISPONIVEL:
        raise HTTPException(500, "LibreOffice não disponível")
    
    try:
        revisoes_list = json.loads(revisoes)
        revisoes_parsed = [RevisaoLibreOffice(**r) for r in revisoes_list]
    except Exception as e:
        raise HTTPException(400, f"JSON inválido: {e}")
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        content = await arquivo.read()
        tmp.write(content)
        input_path = tmp.name
    
    output_path = input_path.replace(".docx", "_REVISADO.docx")
    
    try:
        resultado = _aplicar_revisoes_lo(input_path, revisoes_parsed, autor, output_path)
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total": str(resultado["total"]),
                "X-OK": str(resultado["ok"]),
                "X-Falhas": str(resultado["falhas"])
            }
        )
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@app.post("/libreoffice/aplicar-revisoes-json")
async def libreoffice_aplicar_revisoes_json(
    docx_url: str = Form(...),
    revisoes: str = Form(...),
    autor: str = Form("Revisor IA")
):
    """Aplica revisões via URL do documento (ideal para n8n)."""
    if not LIBREOFFICE_DISPONIVEL:
        raise HTTPException(500, "LibreOffice não disponível")
    
    async with httpx.AsyncClient() as client:
        resp = await client.get(docx_url)
        if resp.status_code != 200:
            raise HTTPException(400, f"Erro ao baixar: {resp.status_code}")
    
    try:
        revisoes_list = json.loads(revisoes)
        revisoes_parsed = [RevisaoLibreOffice(**r) for r in revisoes_list]
    except Exception as e:
        raise HTTPException(400, f"JSON inválido: {e}")
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(resp.content)
        input_path = tmp.name
    
    output_path = input_path.replace(".docx", "_REVISADO.docx")
    
    try:
        resultado = _aplicar_revisoes_lo(input_path, revisoes_parsed, autor, output_path)
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx"
        )
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@app.post("/libreoffice/reset")
async def libreoffice_reset():
    """Reset da conexão com LibreOffice."""
    LibreOfficeConnection.reset()
    return {"msg": "Conexão resetada. Próxima chamada reconectará."}


# ============================================================================
# ENDPOINTS - REVISAO COM TRACK CHANGES (OOXML)
# ============================================================================

@app.post("/revisao/extrair-texto")
async def revisao_extrair_texto(payload: ExtrairTextoDocxPayload):
    """
    Extrai texto estruturado de um documento DOCX para analise.

    Retorna paragrafos com indices para referencia nas revisoes.
    """
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        paragrafos = []
        texto_parts = []
        idx = 0

        for para in doc.paragraphs:
            texto = para.text.strip()
            if not texto:
                continue

            # Detecta tipo do paragrafo
            estilo = para.style.name if para.style else "Normal"
            tipo = "paragraph"
            if "heading" in estilo.lower():
                tipo = estilo.lower().replace(" ", "")

            paragrafos.append({
                "indice": idx,
                "texto": texto,
                "tipo": tipo,
                "tamanho": len(texto)
            })

            texto_parts.append(f"[P{idx}|{tipo.upper()}] {texto}")
            idx += 1

        return {
            "paragrafos": paragrafos,
            "texto_completo": "\n\n".join(texto_parts),
            "total_paragrafos": idx
        }
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/aplicar")
async def revisao_aplicar(payload: AplicarRevisoesPayload):
    """
    Aplica revisoes a um documento DOCX com Track Changes.

    Recebe:
    - docx_url: URL do documento original (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - revisoes: Lista de revisoes no formato:
        {
            "tipo": "SEO|TECNICO|TEXTO",
            "acao": "substituir|deletar|inserir|comentario",
            "texto_original": "texto exato a encontrar",
            "texto_novo": "texto substituto",
            "justificativa": "explicacao"
        }
    - autor: Nome do autor das revisoes

    Retorna: Documento DOCX com Track Changes aplicados
    """
    # Obtém bytes do documento
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    # Salva em arquivo temporario
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_revisado.docx")

    try:
        # Converte revisoes para lista de dicts
        revisoes_list = [r.model_dump() for r in payload.revisoes]

        # Aplica revisoes
        resultado = aplicar_revisoes_docx(
            input_path,
            output_path,
            revisoes_list,
            payload.autor
        )

        # Retorna arquivo com headers de estatisticas
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total-Revisoes": str(resultado["total_revisoes"]),
                "X-Aplicadas": str(resultado["aplicadas"]),
                "X-Falhas": str(resultado["falhas"]),
                "X-Comentarios": str(resultado["comentarios"])
            }
        )
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar revisoes: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@app.post("/revisao/aplicar-json")
async def revisao_aplicar_json(
    docx_url: str = Form(...),
    revisoes: str = Form(...),
    autor: str = Form("Agente IA Revisor")
):
    """
    Aplica revisoes via Form (compativel com n8n HTTP Request).

    Campos Form:
    - docx_url: URL do documento DOCX
    - revisoes: JSON string com array de revisoes
    - autor: Nome do autor (opcional)
    """
    # Baixa o documento
    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.get(docx_url)
        if resp.status_code != 200:
            raise HTTPException(400, f"Erro ao baixar documento: {resp.status_code}")

    # Parse das revisoes
    try:
        revisoes_list = json.loads(revisoes)
    except json.JSONDecodeError as e:
        raise HTTPException(400, f"JSON invalido: {e}")

    # Salva em arquivo temporario
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(resp.content)
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_revisado.docx")

    try:
        # Aplica revisoes
        resultado = aplicar_revisoes_docx(
            input_path,
            output_path,
            revisoes_list,
            autor
        )

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total-Revisoes": str(resultado["total_revisoes"]),
                "X-Aplicadas": str(resultado["aplicadas"]),
                "X-Falhas": str(resultado["falhas"]),
                "X-Comentarios": str(resultado["comentarios"])
            }
        )
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar revisoes: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


# ============================================================================
# ENDPOINTS - AGENTES DE REVISAO (LLM)
# ============================================================================

def _is_image_caption(para) -> bool:
    """
    Detecta se um paragrafo e uma legenda de imagem.
    Legendas sao geradas com: alinhamento centralizado, italico, Pt(10), cor cinza (102,102,102).
    """
    # Verifica alinhamento centralizado
    if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return False

    # Verifica se tem runs
    runs = para.runs
    if not runs:
        return False

    # Verifica se TODOS os runs sao italicos e tem font size 10pt
    for run in runs:
        if not run.italic:
            return False
        if run.font.size and run.font.size != Pt(10):
            return False
        # Verifica cor cinza (102, 102, 102) se definida
        if run.font.color.rgb:
            if run.font.color.rgb != RGBColor(102, 102, 102):
                return False

    return True


def _extrair_texto_para_revisao(docx_path: str, incluir_legendas: bool = False) -> tuple:
    """
    Extrai texto estruturado de um DOCX para analise.

    Args:
        docx_path: Caminho do arquivo DOCX
        incluir_legendas: Se True, inclui legendas de imagem (para agente de imagem)

    Returns:
        (texto_completo, titulo)
    """
    doc = Document(docx_path)
    paragrafos = []
    texto_parts = []
    titulo = ""
    idx = 0

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        # Detecta legendas de imagem
        if _is_image_caption(para):
            if incluir_legendas:
                # Para o agente de imagem, marca como legenda
                paragrafos.append({
                    "indice": idx,
                    "texto": texto,
                    "tipo": "image_caption"
                })
                texto_parts.append(f"[P{idx}|IMAGE_CAPTION] {texto}")
                idx += 1
            # Se nao incluir legendas, simplesmente pula
            continue

        estilo = para.style.name if para.style else "Normal"
        tipo = "paragraph"
        if "heading" in estilo.lower():
            tipo = estilo.lower().replace(" ", "")
            if not titulo and "1" in tipo:
                titulo = texto

        paragrafos.append({
            "indice": idx,
            "texto": texto,
            "tipo": tipo
        })
        texto_parts.append(f"[P{idx}|{tipo.upper()}] {texto}")
        idx += 1

    texto_completo = "\n\n".join(texto_parts)
    return texto_completo, titulo


@app.post("/revisao/agente-seo")
async def revisao_agente_seo(payload: RevisaoAgentPayload):
    """
    Executa o agente de revisao SEO.

    Analisa o documento e retorna JSON com sugestoes de SEO.

    Parametros:
    - docx_url: URL do documento DOCX (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - guia_seo_url: URL do guia de SEO (opcional)
    - url_artigo: URL original do artigo (contexto)
    - titulo: Titulo do artigo (contexto)
    """
    # Debug: log o que está chegando
    print(f"[DEBUG] docx_url: {payload.docx_url[:100] if payload.docx_url else None}...")
    print(f"[DEBUG] docx_base64 type: {type(payload.docx_base64)}")
    print(f"[DEBUG] docx_base64 len: {len(payload.docx_base64) if payload.docx_base64 else 0}")
    print(f"[DEBUG] docx_base64 preview: {payload.docx_base64[:200] if payload.docx_base64 else None}...")

    # Obtém bytes do documento
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
        print(f"[DEBUG] docx_bytes len: {len(docx_bytes)}")
        print(f"[DEBUG] docx_bytes header: {docx_bytes[:20]}")
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    # Baixa guia de SEO se fornecido
    guia_seo = "Use boas praticas gerais de SEO para conteudo tecnico educacional."
    if payload.guia_seo_url:
        try:
            async with httpx.AsyncClient(timeout=60.0) as http_client:
                guia_resp = await http_client.get(payload.guia_seo_url)
                if guia_resp.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp.write(guia_resp.content)
                        guia_path = tmp.name
                    guia_doc = Document(guia_path)
                    guia_seo = "\n".join([p.text for p in guia_doc.paragraphs if p.text.strip()])
                    os.unlink(guia_path)
        except Exception as e:
            print(f"Aviso: Nao foi possivel carregar guia SEO: {e}")

    # Salva documento temporario
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        # Verifica se arquivo existe e tem conteúdo
        if not os.path.exists(tmp_path):
            raise HTTPException(500, f"Arquivo temporario nao foi criado: {tmp_path}")
        file_size = os.path.getsize(tmp_path)
        if file_size == 0:
            raise HTTPException(400, "Documento DOCX vazio ou base64 invalido")

        # Extrai texto
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo = payload.titulo or titulo_extraido

        # Formata prompts
        system_prompt, user_prompt = formatar_prompt_seo(
            conteudo=conteudo,
            titulo=titulo,
            url=payload.url_artigo or "",
            guia_seo=guia_seo
        )

        # Chama LLM
        llm_client = criar_cliente_llm(provider=payload.provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt)
        revisoes = llm_client.extrair_json(resposta)

        # Garante tipo SEO
        for rev in revisoes:
            rev["tipo"] = "SEO"

        return {
            "tipo": "SEO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente SEO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/agente-tecnico")
async def revisao_agente_tecnico(payload: RevisaoAgentPayload):
    """
    Executa o agente de revisao TECNICA.

    Analisa o documento e retorna JSON com correcoes tecnicas.

    Parametros:
    - docx_url: URL do documento DOCX (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - url_artigo: URL original do artigo (contexto)
    - titulo: Titulo do artigo (contexto)
    - data_publicacao: Data de publicacao original
    """
    # Obtém bytes do documento
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    # Salva documento temporario
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        # Verifica se arquivo existe e tem conteúdo
        if not os.path.exists(tmp_path):
            raise HTTPException(500, f"Arquivo temporario nao foi criado: {tmp_path}")
        file_size = os.path.getsize(tmp_path)
        if file_size == 0:
            raise HTTPException(400, "Documento DOCX vazio ou base64 invalido")

        # Extrai texto
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo = payload.titulo or titulo_extraido

        # Formata prompts
        system_prompt, user_prompt = formatar_prompt_tecnico(
            conteudo=conteudo,
            titulo=titulo,
            url=payload.url_artigo or "",
            data_publicacao=payload.data_publicacao or ""
        )

        # Chama LLM (com busca web quando disponivel)
        llm_client = criar_cliente_llm(provider=payload.provider)
        resposta = llm_client.gerar_resposta_com_busca(system_prompt, user_prompt)
        revisoes = llm_client.extrair_json(resposta)

        # Garante tipo TECNICO
        for rev in revisoes:
            rev["tipo"] = "TECNICO"

        return {
            "tipo": "TECNICO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TECNICO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/agente-texto")
async def revisao_agente_texto(payload: RevisaoAgentPayload):
    """
    Executa o agente de revisao TEXTUAL.

    Analisa o documento e retorna JSON com melhorias de texto.

    Parametros:
    - docx_url: URL do documento DOCX (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - url_artigo: URL original do artigo (contexto)
    - titulo: Titulo do artigo (contexto)
    """
    # Obtém bytes do documento
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    # Salva documento temporario
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        # Verifica se arquivo existe e tem conteúdo
        if not os.path.exists(tmp_path):
            raise HTTPException(500, f"Arquivo temporario nao foi criado: {tmp_path}")
        file_size = os.path.getsize(tmp_path)
        if file_size == 0:
            raise HTTPException(400, "Documento DOCX vazio ou base64 invalido")

        # Extrai texto
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo = payload.titulo or titulo_extraido

        # Formata prompts
        system_prompt, user_prompt = formatar_prompt_texto(
            conteudo=conteudo,
            titulo=titulo,
            url=payload.url_artigo or ""
        )

        # Chama LLM
        llm_client = criar_cliente_llm(provider=payload.provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt)
        revisoes = llm_client.extrair_json(resposta)

        # Garante tipo TEXTO
        for rev in revisoes:
            rev["tipo"] = "TEXTO"

        return {
            "tipo": "TEXTO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TEXTO: {str(e)}")
    finally:
        os.unlink(tmp_path)


# ============================================================================
# ENDPOINTS - REVISAO VIA FORM (MULTIPART) - Para uso com n8n
# ============================================================================

@app.post("/revisao/agente-seo-form")
async def revisao_agente_seo_form(
    file: UploadFile = File(...),
    provider: str = Form("anthropic"),
    url_artigo: str = Form(""),
    titulo: str = Form(""),
    guia_seo_file: Optional[UploadFile] = File(None),
    palavras_chave: str = Form("")
):
    """
    Executa o agente de revisao SEO via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.

    Parametros:
    - file: Arquivo DOCX do artigo a ser revisado
    - provider: "anthropic" ou "openai"
    - url_artigo: URL original do artigo
    - titulo: Titulo do artigo (opcional, extraido do DOCX se nao fornecido)
    - guia_seo_file: Arquivo DOCX com o guia de SEO da empresa (opcional)
    - palavras_chave: Palavras-chave do Google separadas por virgula ou quebra de linha
    """
    # Salva arquivo temporario
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    guia_path = None

    try:
        # Extrai texto do artigo
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        # Extrai guia SEO se fornecido
        guia_seo = "Use boas praticas gerais de SEO para conteudo tecnico educacional."
        if guia_seo_file and guia_seo_file.filename:
            guia_bytes = await guia_seo_file.read()
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_guia:
                tmp_guia.write(guia_bytes)
                tmp_guia.flush()
                guia_path = tmp_guia.name

            guia_doc = Document(guia_path)
            guia_seo = "\n".join([p.text for p in guia_doc.paragraphs if p.text.strip()])

        # Formata palavras-chave
        palavras_formatadas = "Nenhuma palavra-chave especifica fornecida. Use seu conhecimento de SEO."
        if palavras_chave and palavras_chave.strip():
            # Aceita separacao por virgula ou quebra de linha
            keywords = [kw.strip() for kw in palavras_chave.replace('\n', ',').split(',') if kw.strip()]
            if keywords:
                palavras_formatadas = "\n".join([f"- {kw}" for kw in keywords])

        # Formata prompts
        system_prompt, user_prompt = formatar_prompt_seo(
            conteudo=conteudo,
            titulo=titulo_final,
            url=url_artigo,
            guia_seo=guia_seo,
            palavras_chave=palavras_formatadas
        )

        # Chama LLM
        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "SEO"

        return {
            "tipo": "SEO",
            "total_sugestoes": len(revisoes),
            "palavras_chave_usadas": palavras_chave.strip() if palavras_chave else None,
            "guia_seo_fornecido": bool(guia_seo_file and guia_seo_file.filename),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente SEO: {str(e)}")
    finally:
        os.unlink(tmp_path)
        if guia_path and os.path.exists(guia_path):
            os.unlink(guia_path)


@app.post("/revisao/agente-tecnico-form")
async def revisao_agente_tecnico_form(
    file: UploadFile = File(...),
    provider: str = Form("anthropic"),
    url_artigo: str = Form(""),
    titulo: str = Form(""),
    data_publicacao: str = Form("")
):
    """
    Executa o agente de revisao TECNICA via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.
    """
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        system_prompt, user_prompt = formatar_prompt_tecnico(
            conteudo=conteudo,
            titulo=titulo_final,
            url=url_artigo,
            data_publicacao=data_publicacao
        )

        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta_com_busca(system_prompt, user_prompt)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TECNICO"

        return {
            "tipo": "TECNICO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TECNICO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/agente-texto-form")
async def revisao_agente_texto_form(
    file: UploadFile = File(...),
    provider: str = Form("anthropic"),
    url_artigo: str = Form(""),
    titulo: str = Form("")
):
    """
    Executa o agente de revisao TEXTUAL via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.
    """
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        system_prompt, user_prompt = formatar_prompt_texto(
            conteudo=conteudo,
            titulo=titulo_final,
            url=url_artigo
        )

        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TEXTO"

        return {
            "tipo": "TEXTO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TEXTO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/agente-imagem")
async def revisao_agente_imagem(payload: RevisaoImagemPayload):
    """
    Executa o agente de revisao de IMAGENS.

    Analisa as imagens do artigo quanto a relevancia, qualidade,
    atualizacao de screenshots e acessibilidade (alt text).

    Parametros:
    - docx_url: URL do documento DOCX (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - url_artigo: URL original do artigo (OBRIGATORIO para extrair imagens)
    - provider: "anthropic" ou "openai" (default: anthropic)
    - titulo: Titulo do artigo (opcional)

    Retorna:
    - JSON com sugestoes de revisao de imagens
    """
    from datetime import datetime

    if not payload.url_artigo:
        raise HTTPException(400, "url_artigo e obrigatorio para o agente de imagem")

    try:
        # Obtem bytes do DOCX
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)

        # Salva temporariamente
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(docx_bytes)
            tmp.flush()
            tmp_path = tmp.name

        try:
            # Extrai texto do DOCX (incluindo legendas para o agente de imagem)
            conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path, incluir_legendas=True)
            titulo_final = payload.titulo or titulo_extraido

            # Faz scraping do artigo para obter as imagens
            print(f"📥 Extraindo imagens de: {payload.url_artigo}")
            with httpx.Client(timeout=30, follow_redirects=True) as http_client:
                response = http_client.get(payload.url_artigo)
                response.raise_for_status()
                html = response.text

            article_data = extract_article_content(html, payload.url_artigo)

            # Filtra apenas itens do tipo "image"
            imagens = [item for item in article_data.get('content', []) if item.get('type') == 'image']
            print(f"📸 {len(imagens)} imagens encontradas")

            if not imagens:
                return {
                    "tipo": "IMAGEM",
                    "total_sugestoes": 0,
                    "revisoes": [],
                    "mensagem": "Nenhuma imagem encontrada no artigo"
                }

            # Prepara prompt
            data_atual = datetime.now().strftime("%d/%m/%Y")
            system_prompt, user_prompt = formatar_prompt_imagem(
                conteudo=conteudo,
                imagens=imagens,
                titulo=titulo_final,
                url=payload.url_artigo,
                data_atual=data_atual
            )

            # Chama LLM com visao multimodal
            llm_client = criar_cliente_llm(provider=payload.provider)

            # Anthropic: visao + busca web | OpenAI: apenas visao
            resposta = llm_client.gerar_resposta_com_imagens_e_busca(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                imagens=imagens
            )

            revisoes = llm_client.extrair_json(resposta)

            # Garante tipo IMAGEM em todas as revisoes
            for rev in revisoes:
                rev["tipo"] = "IMAGEM"

            return {
                "tipo": "IMAGEM",
                "total_sugestoes": len(revisoes),
                "total_imagens": len(imagens),
                "revisoes": revisoes
            }

        finally:
            os.unlink(tmp_path)

    except httpx.HTTPError as e:
        raise HTTPException(400, f"Erro ao buscar URL do artigo: {str(e)}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Erro no agente IMAGEM: {str(e)}")


@app.post("/revisao/agente-imagem-form")
async def revisao_agente_imagem_form(
    file: UploadFile = File(...),
    url_artigo: str = Form(...),
    provider: str = Form("anthropic"),
    titulo: str = Form("")
):
    """
    Executa o agente de revisao de IMAGENS via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.

    Parametros:
    - file: Arquivo DOCX
    - url_artigo: URL original do artigo (OBRIGATORIO para extrair imagens)
    - provider: "anthropic" ou "openai" (default: anthropic)
    - titulo: Titulo do artigo (opcional)
    """
    from datetime import datetime

    if not url_artigo:
        raise HTTPException(400, "url_artigo e obrigatorio para o agente de imagem")

    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        # Extrai texto do DOCX (incluindo legendas para o agente de imagem)
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path, incluir_legendas=True)
        titulo_final = titulo or titulo_extraido

        # Faz scraping do artigo para obter as imagens
        print(f"📥 Extraindo imagens de: {url_artigo}")
        with httpx.Client(timeout=30, follow_redirects=True) as http_client:
            response = http_client.get(url_artigo)
            response.raise_for_status()
            html = response.text

        article_data = extract_article_content(html, url_artigo)

        # Filtra apenas itens do tipo "image"
        imagens = [item for item in article_data.get('content', []) if item.get('type') == 'image']
        print(f"📸 {len(imagens)} imagens encontradas")

        if not imagens:
            return {
                "tipo": "IMAGEM",
                "total_sugestoes": 0,
                "revisoes": [],
                "mensagem": "Nenhuma imagem encontrada no artigo"
            }

        # Prepara prompt
        data_atual = datetime.now().strftime("%d/%m/%Y")
        system_prompt, user_prompt = formatar_prompt_imagem(
            conteudo=conteudo,
            imagens=imagens,
            titulo=titulo_final,
            url=url_artigo,
            data_atual=data_atual
        )

        # Chama LLM com visao multimodal
        llm_client = criar_cliente_llm(provider=provider)

        # Anthropic: visao + busca web | OpenAI: apenas visao
        resposta = llm_client.gerar_resposta_com_imagens_e_busca(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            imagens=imagens
        )

        revisoes = llm_client.extrair_json(resposta)

        # Garante tipo IMAGEM em todas as revisoes
        for rev in revisoes:
            rev["tipo"] = "IMAGEM"

        return {
            "tipo": "IMAGEM",
            "total_sugestoes": len(revisoes),
            "total_imagens": len(imagens),
            "revisoes": revisoes
        }

    except httpx.HTTPError as e:
        raise HTTPException(400, f"Erro ao buscar URL do artigo: {str(e)}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Erro no agente IMAGEM: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/aplicar-form")
async def revisao_aplicar_form(
    file: UploadFile = File(...),
    revisoes: str = Form(...),
    autor: str = Form("Agente IA Revisor")
):
    """
    Aplica revisoes a um documento DOCX via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.

    - file: Arquivo DOCX
    - revisoes: JSON string com lista de revisoes
    - autor: Nome do autor das revisoes
    """
    import json as json_lib

    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_revisado.docx")

    try:
        revisoes_list = json_lib.loads(revisoes)

        resultado = aplicar_revisoes_docx(
            input_path,
            output_path,
            revisoes_list,
            autor
        )

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total-Revisoes": str(resultado["total_revisoes"]),
                "X-Aplicadas": str(resultado["aplicadas"]),
                "X-Falhas": str(resultado["falhas"]),
                "X-Comentarios": str(resultado["comentarios"])
            }
        )
    except json_lib.JSONDecodeError as e:
        raise HTTPException(400, f"JSON de revisoes invalido: {str(e)}")
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar revisoes: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@app.post("/revisao/aplicar-comentarios-form")
async def revisao_aplicar_comentarios_form(
    file: UploadFile = File(...),
    revisoes: str = Form(...),
    autor: str = Form("Agente IA Revisor")
):
    """
    Aplica SOMENTE comentarios a um documento DOCX via upload de arquivo.
    Nao altera o texto do documento - apenas adiciona comentarios nos trechos encontrados.
    Suporta multiplos comentarios no mesmo trecho (ranges sobrepostos).

    - file: Arquivo DOCX
    - revisoes: JSON string com lista de revisoes
    - autor: Nome do autor dos comentarios
    """
    import json as json_lib

    # Preserva o nome original do arquivo
    original_filename = file.filename or "documento.docx"

    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_comentado.docx")

    try:
        revisoes_list = json_lib.loads(revisoes)

        resultado = aplicar_comentarios_docx(
            input_path,
            output_path,
            revisoes_list,
            autor
        )

        stats = resultado.get('estatisticas', {})

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=original_filename,
            headers={
                "X-Total-Comentarios": str(resultado.get('total_comentarios', 0)),
                "X-Match-Exato": str(stats.get('exato', 0)),
                "X-Match-Normalizado": str(stats.get('normalizado', 0)),
                "X-Match-Fuzzy": str(stats.get('fuzzy', 0)),
                "X-Match-Paragrafo": str(stats.get('paragrafo', 0)),
                "X-Match-Falhas": str(stats.get('falha', 0)),
            }
        )
    except json_lib.JSONDecodeError as e:
        raise HTTPException(400, f"JSON de revisoes invalido: {str(e)}")
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar comentarios: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


# ============================================================================
# ENDPOINTS - VÍDEO
# ============================================================================

@app.post("/processar_video_urls")
async def processar_video_urls(payload: VideoURLProcessingPayload, background_tasks: BackgroundTasks):
    job_id = str(uuid.uuid4())
    job_dir = TEMP_DIR / job_id
    job_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        video_paths = []
        for i, url in enumerate(payload.video_urls):
            video_path = job_dir / f"video_{i:03d}.mp4"
            baixar_arquivo(url, str(video_path))
            video_paths.append(str(video_path))
        
        audio_path = job_dir / "audio_narracao.mp3"
        baixar_arquivo(payload.audio_url, str(audio_path))
        
        srt_path = None
        if payload.adicionar_legendas:
            srt_path = str(job_dir / "legendas.srt")
            gerar_legendas_srt(str(audio_path), srt_path)
        
        output_path = job_dir / "video_final.mp4"
        
        criar_video_com_transicoes(
            video_paths, str(audio_path), str(output_path),
            transicao_duracao=payload.transicao_duracao,
            transicao_tipo=payload.transicao_tipo,
            legendas_srt=srt_path,
            estilo_legenda=payload.estilo_legenda,
            legenda_config=payload.legenda_config
        )
        
        background_tasks.add_task(cleanup_job, job_dir, 3600)
        
        filename = payload.output_filename if payload.output_filename.endswith('.mp4') else f"{payload.output_filename}.mp4"
        return FileResponse(
            path=str(output_path),
            media_type="video/mp4",
            filename=filename,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    
    except Exception as e:
        if job_dir.exists():
            shutil.rmtree(job_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"Erro ao processar vídeo: {str(e)}")


@app.post("/processar_video")
async def processar_video(
    background_tasks: BackgroundTasks,
    videos: List[UploadFile] = File(...),
    audio: UploadFile = File(...),
    transicao_duracao: float = 0.5,
    transicao_tipo: str = "fade"
):
    job_id = str(uuid.uuid4())
    job_dir = TEMP_DIR / job_id
    job_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        video_paths = []
        for i, video in enumerate(videos):
            video_path = job_dir / f"video_{i:03d}.mp4"
            with open(video_path, "wb") as f:
                shutil.copyfileobj(video.file, f)
            video_paths.append(str(video_path))
        
        audio_path = job_dir / "audio_narracao.mp3"
        with open(audio_path, "wb") as f:
            shutil.copyfileobj(audio.file, f)
        
        output_path = job_dir / "video_final.mp4"
        
        criar_video_com_transicoes(
            video_paths, str(audio_path), str(output_path),
            transicao_duracao=transicao_duracao,
            transicao_tipo=transicao_tipo
        )
        
        background_tasks.add_task(cleanup_job, job_dir, 3600)
        
        return FileResponse(
            path=str(output_path),
            media_type="video/mp4",
            filename=f"video_final_{job_id[:8]}.mp4"
        )
    
    except Exception as e:
        if job_dir.exists():
            shutil.rmtree(job_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"Erro ao processar vídeo: {str(e)}")


@app.get("/processar_video/status")
def status_processamento():
    return {
        "ok": True,
        "ffmpeg_disponivel": shutil.which("ffmpeg") is not None,
        "temp_dir": str(TEMP_DIR),
        "transicoes_disponiveis": [
            "fade", "wipeleft", "wiperight", "wipeup", "wipedown",
            "slideleft", "slideright", "slideup", "slidedown",
            "dissolve", "pixelize"
        ]
    }


# ============================================================================
# ENDPOINTS - LINKEDIN/ALURA
# ============================================================================

@app.post("/pesquisa_mercado_linkedin")
def pesquisa_mercado_linkedin(p: PesquisaPayload):
    params = {"keywords": p.query, "location": "Brasil", "start": 0}
    user = os.environ.get("LINKEDIN_USER")
    passwd = os.environ.get("LINKEDIN_PASS")
    if not user or not passwd:
        raise HTTPException(status_code=500, detail="Defina LINKEDIN_USER e LINKEDIN_PASS")
    
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True, args=["--no-sandbox", "--disable-setuid-sandbox"])
            page = browser.new_page()
            login_linkedin(page, user, passwd)
            links = []
            for i in range(0, int(p.n_vagas), 25):
                params["start"] = i
                page.goto(f"https://www.linkedin.com/jobs/search/?{urlencode(params)}")
                lista = page.locator("div.scaffold-layout__list")
                lista.first.wait_for(state="visible", timeout=60000)
                results = page.locator("div.jobs-search-results-list").first
                container = results if results.count() > 0 else lista.first
                page.wait_for_selector('a[href^="/jobs/view/"]', timeout=60000)
                vagas = rolar_e_coletar_vagas(page, container, max_rolagens=10, pausa=1.2)
                links = list(dict.fromkeys(links + vagas))
            page.goto("https://www.linkedin.com/m/logout/")
            browser.close()
        return Response(
            content=json.dumps({"ok": True, "data": links}, ensure_ascii=False),
            media_type="application/json"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Falha: {e}")


@app.post("/cadastrar_curso")
def cadastrar(p: Payload):
    instrutores_path = "/files/data/instrutores.json"
    if not os.path.exists(instrutores_path):
        raise HTTPException(status_code=500, detail=f"Arquivo não encontrado: {instrutores_path}")
    with open(instrutores_path, "r", encoding="utf-8") as f:
        instrutores = json.load(f)
    autor_valor = next((a["valor"] for a in instrutores if a["nome"] == p.nome_instrutor), None)
    if not autor_valor:
        raise HTTPException(status_code=404, detail="Instrutor não localizado.")
    user = os.environ.get("ALURA_USER")
    passwd = os.environ.get("ALURA_PASS")
    if not user or not passwd:
        raise HTTPException(status_code=500, detail="Defina ALURA_USER e ALURA_PASS")
    code = gerar_codigo_cursos(p.nome_curso)
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            page = browser.new_page()
            login_alura(page, user, passwd)
            page.goto("https://cursos.alura.com.br/admin/v2/newCourse")
            page.fill('input[name="name"]', p.nome_curso)
            page.fill('input[name="code"]', code)
            page.fill('input[name="estimatedTimeToFinish"]', str(int(p.tempo_curso)))
            page.fill('input[name="metadescription"]', 'Será atualizado pelo(a) instrutor(a).')
            page.select_option('select[name="authors"]', value=autor_valor)
            browser.close()
        return Response(
            content=json.dumps({"ok": True, "code": code}, ensure_ascii=False),
            media_type="application/json"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Falha: {e}")


@app.post("/get_transcription_course")
def get_transcription_course(p: IDPayload):
    user = os.environ.get("ALURA_USER")
    passwd = os.environ.get("ALURA_PASS")
    if not user or not passwd:
        raise HTTPException(status_code=500, detail="Defina ALURA_USER e ALURA_PASS")
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            page = browser.new_page()
            login_alura(page, user, passwd)

            page.goto(f"https://cursos.alura.com.br/admin/courses/v2/{p.id}", timeout=60000, wait_until="networkidle")
            page.wait_for_selector('div.form-group', timeout=60000)
            link_href = page.evaluate('''() => {
                const links = document.querySelectorAll('a.btn-default');
                for (let link of links) {
                    if (link.href.includes('/course/') && link.textContent.includes('Ver curso')) {
                        return link.getAttribute('href');
                    }
                }
                return null;
            }''')

            if not link_href:
                raise Exception("Não achou o link 'Ver curso'")
            link = "https://cursos.alura.com.br" + link_href

            page.goto(link, timeout=60000, wait_until="domcontentloaded")
            page.wait_for_selector(".courseSectionList", timeout=60000)
            html = page.content()
            soup = BeautifulSoup(html, "html.parser")
            nome = soup.find("h1").strong.get_text()
            videos = []
            for item in soup.find_all("li", class_="courseSection-listItem"):
                aula = f"https://cursos.alura.com.br{item.find('a', class_='courseSectionList-section')['href']}"
                page.goto(aula, timeout=60000, wait_until="domcontentloaded")
                page.wait_for_selector(".task-menu-sections-select", timeout=60000)
                html = page.content()
                soup_section = BeautifulSoup(html, "html.parser")
                for video in soup_section.find_all("a", class_="task-menu-nav-item-link task-menu-nav-item-link-VIDEO"):
                    videos.append(f"https://cursos.alura.com.br{video['href']}")
            transcricoes = []
            for index, video in enumerate(videos):
                page.goto(video, timeout=60000, wait_until="domcontentloaded")
                page.wait_for_selector("#transcription", timeout=60000)
                html = page.content()
                soup_video = BeautifulSoup(html, "html.parser")
                title = soup_video.find("h1", class_="task-body-header-title").span.get_text()
                transcription = soup_video.find("section", id="transcription").get_text()
                transcription = transcription.replace("Transcrição", f"Vídeo {index + 1} -{title}")
                transcricoes.append(limpar_texto(transcription))
            browser.close()
        return Response(
            content=json.dumps({
                "id": p.id,
                "nome": nome,
                "link": link,
                "transcricao": transcricoes
            }, ensure_ascii=False),
            media_type="application/json"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Falha: {e}")